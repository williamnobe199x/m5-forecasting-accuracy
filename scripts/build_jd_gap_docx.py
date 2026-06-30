"""Create a Word brief for JD requirement gaps in the M5 project."""

from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--job-dir", default=ROOT / "output/job_alignment", type=Path)
    parser.add_argument("--project-summary", default=ROOT / "output/project_summary.json", type=Path)
    parser.add_argument("--out", default=ROOT / "docs/M5_JD_gap_analysis.docx", type=Path)
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("M5 Accuracy 项目 JD 能力缺口分析")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    text = "基于 9 张预测/供应链/数据科学岗位截图整理，用于成果汇报和后续项目补强。"
    run = subtitle.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("667085")


def add_callout(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    set_paragraph_shading(p, CALLOUT_FILL)
    label_run = p.add_run(label + "：")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(body)


def add_matrix_table(doc: Document, jd: pd.DataFrame) -> None:
    columns = [
        ("theme", "能力主题", 1550),
        ("status", "状态", 900),
        ("project_evidence", "项目证据", 3150),
        ("hr_talking_point", "HR 关注点", 2200),
        ("next_action", "下一步", 1560),
    ]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, (_col, label, _width) in enumerate(columns):
        table.cell(0, idx).text = label
        set_cell_shading(table.cell(0, idx), LIGHT_FILL)
    repeat_table_header(table.rows[0])

    for _, item in jd.iterrows():
        cells = table.add_row().cells
        for idx, (col, _label, _width) in enumerate(columns):
            cells[idx].text = str(item[col])
    set_table_width(table, [width for _col, _label, width in columns])


def add_gap_table(doc: Document, gaps: pd.DataFrame) -> None:
    columns = [
        ("priority", "优先级", 760),
        ("gap_area", "缺口", 1380),
        ("jd_requirement", "JD 要求", 2820),
        ("current_project_boundary", "当前边界", 2380),
        ("next_action", "补齐动作", 2020),
    ]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, (_col, label, _width) in enumerate(columns):
        table.cell(0, idx).text = label
        set_cell_shading(table.cell(0, idx), LIGHT_FILL)
    repeat_table_header(table.rows[0])

    for _, item in gaps.iterrows():
        cells = table.add_row().cells
        for idx, (col, _label, _width) in enumerate(columns):
            cells[idx].text = str(item[col])
    set_table_width(table, [width for _col, _label, width in columns])


def add_artifact_table(doc: Document) -> None:
    rows = [
        ("网页汇报", "docs/index.html", "成果展示、讲稿、图表和岗位对齐章节"),
        ("讲稿 Markdown", "docs/presentation_script.md", "与网页逐段匹配，可直接排练"),
        ("JD 能力矩阵", "output/job_alignment/jd_capability_matrix.csv", "覆盖/部分覆盖/未覆盖状态"),
        ("JD 缺口清单", "output/job_alignment/jd_gap_list.csv", "按 P0-P3 排序的补齐路线"),
        ("Rolling CV", "output/job_alignment/rolling_cv_summary.csv", "多窗口 WAPE/Bias 验证"),
        ("Direct 试点", "output/direct_segment_lightgbm_wrmsse_by_level.csv", "4 段 direct horizon LightGBM 评分"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["产物", "路径", "用途"]
    for idx, label in enumerate(headers):
        table.cell(0, idx).text = label
        set_cell_shading(table.cell(0, idx), LIGHT_FILL)
    repeat_table_header(table.rows[0])
    for name, path, use in rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = path
        cells[2].text = use
    set_table_width(table, [1500, 3550, 4310])


def add_route(doc: Document) -> None:
    doc.add_heading("三、建议补齐路线", level=1)
    routes = [
        ("P0", "把 item、dept-store、cat-store 聚合 lag/rolling 接入主训练，并用 rolling CV 重新比较。"),
        ("P0", "重做 direct horizon LightGBM：按 F1-F7、F8-F14、F15-F21、F22-F28 训练，加入早停、特征筛选和 ensemble。"),
        ("P1", "补 DuckDB/SQL 指标口径层，展示复杂关联、窗口函数、数据质量校验和可复用数据集市。"),
        ("P1", "增加库存补货仿真：安全库存、reorder point、服务水平、库存周转和订单满足率。"),
        ("P2", "增加促销/节假日 what-if 或准实验设计，形成策略评估而不是只做预测。"),
        ("P2", "建立小型运筹优化 demo，把预测需求转成补货或生产约束下的决策方案。"),
        ("P3", "在预测链路稳定后，再扩展 RAG 复盘助手或 foundation model ensemble。"),
    ]
    for label, text in routes:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(label + "：").bold = True
        p.add_run(text)


def add_hr_talking_points(doc: Document) -> None:
    doc.add_heading("四、成果汇报中建议强调的 HR 关注点", level=1)
    points = [
        "不是只追单个分数，而是有数据处理、特征、模型、评分、可视化、复盘的完整链路。",
        "能用 WRMSSE、WAPE、Bias、rolling CV 等指标说明预测质量和稳定性。",
        "能解释为什么当前不盲目换大模型：基础实验体系和多粒度特征更优先。",
        "能诚实说明项目边界，并把缺口转化为下一阶段路线图。",
    ]
    for point in points:
        doc.add_paragraph(point, style="List Bullet")


def main() -> None:
    args = parse_args()
    jd = pd.read_csv(args.job_dir / "jd_capability_matrix.csv")
    gaps = pd.read_csv(args.job_dir / "jd_gap_list.csv")

    doc = Document()
    style_document(doc)
    add_title(doc)
    add_callout(
        doc,
        "结论",
        "当前项目最适合定位为“零售销售预测 + 可解释评估 + 成果汇报”作品；它已经覆盖预测建模、稳定实验、可视化表达等核心能力，但尚未覆盖完整供应链计划、库存补货、SQL 数据治理和系统上线闭环。",
    )

    doc.add_heading("一、岗位能力覆盖矩阵", level=1)
    p = doc.add_paragraph()
    p.add_run("读取来源：").bold = True
    p.add_run("9 张岗位截图被归纳为 12 个能力主题，状态分为 Covered、Partial、Missing。")
    add_matrix_table(doc, jd)

    doc.add_heading("二、未满足或部分满足的要求", level=1)
    doc.add_paragraph(
        "下面的清单适合在面试或成果汇报后半段使用：先说明已经完成的能力证据，再主动展示下一步补齐路线。"
    )
    add_gap_table(doc, gaps)

    add_route(doc)
    add_hr_talking_points(doc)

    doc.add_heading("五、关联产物路径", level=1)
    add_artifact_table(doc)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(f"wrote {args.out.name}")


if __name__ == "__main__":
    main()
